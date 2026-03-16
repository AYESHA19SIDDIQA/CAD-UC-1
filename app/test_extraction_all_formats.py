"""
Test script for extracting text from sample sanction documents
Handles .doc, .docx, and PDF formats
"""
import sys
import os
from pathlib import Path

# Add parent directory to path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

def extract_from_doc_legacy(file_path):
    """Try to extract from legacy .doc file using win32com"""
    try:
        import win32com.client
        
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        doc = word.Documents.Open(os.path.abspath(file_path))
        text = doc.Content.Text
        
        # Extract tables
        tables_text = []
        for table in doc.Tables:
            for row in table.Rows:
                row_text = []
                for cell in row.Cells:
                    row_text.append(cell.Range.Text.strip())
                tables_text.append(" | ".join(row_text))
        
        doc.Close()
        word.Quit()
        
        full_text = text + "\n" + "\n".join(tables_text)
        return full_text
    except Exception as e:
        print(f"Error with win32com extraction: {e}")
        return None

def extract_from_docx(file_path):
    """Extract from .docx file"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        full_text = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error with docx extraction: {e}")
        return None

def convert_doc_to_text_via_textract(file_path):
    """Try using textract library"""
    try:
        import textract
        text = textract.process(file_path).decode('utf-8')
        return text
    except Exception as e:
        print(f"Error with textract: {e}")
        return None

def test_file(file_path):
    """Test extraction on a single file"""
    print("\n" + "="*80)
    print(f"Processing: {os.path.basename(file_path)}")
    print("="*80)
    
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return
    
    file_size = os.path.getsize(file_path)
    print(f"File size: {file_size} bytes")
    
    # Determine file type
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    text = None
    
    # Try different extraction methods
    if ext == '.docx':
        print("Attempting .docx extraction...")
        text = extract_from_docx(file_path)
    elif ext == '.doc':
        print("Attempting legacy .doc extraction with win32com...")
        text = extract_from_doc_legacy(file_path)
    else:
        print(f"Unknown extension: {ext}")
        print("Trying as .doc file...")
        text = extract_from_doc_legacy(file_path)
        
        if not text:
            print("Trying as .docx file...")
            text = extract_from_docx(file_path)
    
    # Fallback to textract
    if not text:
        print("Trying textract library...")
        text = convert_doc_to_text_via_textract(file_path)
    
    if text:
        print("\n--- EXTRACTION SUCCESSFUL ---")
        print(f"Extracted text length: {len(text)} characters")
        print("\n--- FIRST 1000 CHARACTERS ---")
        print(text[:1000])
        print("\n--- LAST 500 CHARACTERS ---")
        print(text[-500:])
        
        # Save to file
        output_file = file_path + "_extracted.txt"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"\nFull text saved to: {output_file}")
    else:
        print("\n❌ EXTRACTION FAILED - Could not extract text from file")
        print("\nTrying to read as binary to check file format...")
        with open(file_path, 'rb') as f:
            header = f.read(8)
            print(f"File header (hex): {header.hex()}")
            print(f"File header (bytes): {header}")

def main():
    """Test extraction on all sample files"""
    
    print("="*80)
    print("TESTING DOCUMENT EXTRACTION")
    print("="*80)
    
    # Check for samples directory
    samples_dir = r"c:\Users\hp\Desktop\BM_stuff\CAD_01\document_generator\app\samples"
    
    if os.path.exists(samples_dir):
        print(f"\nScanning samples directory: {samples_dir}")
        files = list(Path(samples_dir).glob("*"))
        print(f"Found {len(files)} files")
        
        for file_path in files:
            if file_path.is_file():
                test_file(str(file_path))
    else:
        print(f"Samples directory not found: {samples_dir}")
    
    # Also test the main document
    main_doc = r"c:\Users\hp\Desktop\BM_stuff\CAD_01\Sanction Advice Review and Document Generator.docx"
    if os.path.exists(main_doc):
        test_file(main_doc)

if __name__ == "__main__":
    main()
