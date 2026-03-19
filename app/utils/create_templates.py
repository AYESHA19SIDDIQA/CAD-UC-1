"""
Template creator — run this ONCE to generate starter .docx template files
in app/templates/.

Each template contains the correct placeholder tokens.
After running this script, open the templates in Microsoft Word and format
them to your bank's brand guidelines (fonts, logos, headers, footers).
The placeholder tokens {{LIKE_THIS}} will still be found and replaced when
documents are generated — formatting around them is preserved.

Usage:
    python -m app.utils.create_templates
"""
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


TEMPLATE_DIR = Path(__file__).parent.parent / "templates"
TEMPLATE_DIR.mkdir(exist_ok=True)


def _add_header(doc: Document, title: str):
    """Add a simple letterhead block."""
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Date: {{DATE}}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("")


def _add_address_block(doc: Document):
    doc.add_paragraph("To,")
    doc.add_paragraph("{{CUSTOMER_NAME}}")
    doc.add_paragraph("{{CUSTOMER_LOCATION}}")
    doc.add_paragraph("")


def _add_signature(doc: Document):
    doc.add_paragraph("")
    doc.add_paragraph("_" * 35 + "          " + "_" * 35)
    doc.add_paragraph(
        "Authorised Signatory                              Relationship Manager"
    )


def create_offer_letter():
    """Offer Letter — one per sanction (covers all facilities)."""
    doc = Document()
    _add_header(doc, "Offer Letter")
    _add_address_block(doc)

    doc.add_paragraph("Dear Sir / Madam,")
    doc.add_paragraph("")
    doc.add_paragraph(
        "We are pleased to inform you that the following credit facility has been "
        "sanctioned by the competent authority vide Approval No. {{APPROVAL_NO}} "
        "dated {{SANCTION_DATE}}:"
    )
    doc.add_paragraph("")

    # Facility details table
    doc.add_heading("Facility Details", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Details"
    rows = [
        ("Facility Type", "{{FACILITY_TYPE}}"),
        ("Nature of Limit", "{{NATURE_OF_LIMIT}}"),
        ("Approved Limit", "{{CURRENCY}} {{APPROVED_LIMIT}} millions"),
        ("Profit Rate", "{{PROFIT_RATE}}"),
        ("Tenor", "{{TENOR}}"),
        ("Security / Collateral", "{{SECURITY}}"),
        ("Purpose", "{{PURPOSE}}"),
        ("ICRR", "{{ICRR}}"),
        ("Business Segment", "{{BUSINESS_SEGMENT}}"),
    ]
    for label, value in rows:
        r = table.add_row().cells
        r[0].text = label
        r[1].text = value

    doc.add_paragraph("")
    doc.add_paragraph(
        "Kindly confirm your acceptance of the above terms and conditions in writing "
        "within 7 days of the date of this letter."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Yours faithfully,")
    _add_signature(doc)

    path = TEMPLATE_DIR / "offer_letter.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")


def create_sanction_letter():
    doc = Document()
    _add_header(doc, "Sanction Letter")
    _add_address_block(doc)

    doc.add_paragraph(
        "This is to confirm that the following credit facility has been approved for "
        "{{CUSTOMER_NAME}} vide Approval No. {{APPROVAL_NO}} dated {{SANCTION_DATE}}."
    )
    doc.add_paragraph("")
    doc.add_heading("Sanctioned Facility", level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Details"
    for label, value in [
        ("Facility Type", "{{FACILITY_TYPE}}"),
        ("Approved Limit", "{{CURRENCY}} {{APPROVED_LIMIT}} millions"),
        ("Tenor", "{{TENOR}}"),
        ("Profit / Commission", "{{PROFIT_RATE}}"),
        ("Security", "{{SECURITY}}"),
    ]:
        r = table.add_row().cells
        r[0].text = label
        r[1].text = value

    doc.add_paragraph("")
    doc.add_heading("General Conditions", level=1)
    conditions = [
        "Confirmation in writing as to acceptance of approval is obtained from the "
        "customer before allowing facilities.",
        "Ensure proper compliance of Prudential Regulations, bank policy and SBP "
        "guidelines issued from time to time.",
        "In no case withdrawal will be allowed in excess of the approved limit.",
        "All property and charge documents to be completed as per legal opinion "
        "before disbursement.",
        "The Bank reserves the right to add, amend or alter any of the above "
        "conditions at its discretion.",
    ]
    for i, c in enumerate(conditions, 1):
        doc.add_paragraph(f"{i}. {c}")

    _add_signature(doc)
    path = TEMPLATE_DIR / "sanction_letter.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")


def create_terms_conditions_sheet():
    doc = Document()
    _add_header(doc, "Terms and Conditions Sheet")
    _add_address_block(doc)

    doc.add_paragraph(
        "The following terms and conditions apply to the credit facility granted to "
        "{{CUSTOMER_NAME}} (Approval No. {{APPROVAL_NO}}):"
    )
    doc.add_paragraph("")
    doc.add_heading("Facility Summary", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Parameter"
    table.rows[0].cells[1].text = "Detail"
    for l, v in [
        ("Customer", "{{CUSTOMER_NAME}}"),
        ("Facility Type", "{{FACILITY_TYPE}}"),
        ("Approved Limit", "{{CURRENCY}} {{APPROVED_LIMIT}} millions"),
        ("Tenor", "{{TENOR}}"),
        ("Profit Rate", "{{PROFIT_RATE}}"),
        ("Security", "{{SECURITY}}"),
    ]:
        r = table.add_row().cells
        r[0].text = l
        r[1].text = v

    doc.add_paragraph("")
    doc.add_heading("Conditions Precedent to Disbursement", level=1)
    for i, c in enumerate([
        "Execution of all facility documents.",
        "Completion of all security / collateral perfection.",
        "Receipt of all required approvals and NOCs.",
        "KYC documentation up to date.",
    ], 1):
        doc.add_paragraph(f"{i}. {c}")

    _add_signature(doc)
    path = TEMPLATE_DIR / "terms_and_conditions_sheet.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")


def create_demand_promissory_note():
    doc = Document()
    _add_header(doc, "Demand Promissory Note")

    doc.add_paragraph(
        "On demand, I/We, {{CUSTOMER_NAME}}, promise to pay to "
        "Meezan Bank Limited or order, the sum of {{CURRENCY}} {{APPROVED_LIMIT}} "
        "millions (or such part thereof as shall remain outstanding) together with "
        "profit at the rate of {{PROFIT_RATE}} per annum from the date of drawing."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Place of Execution: {{CUSTOMER_LOCATION}}")
    doc.add_paragraph("Date: {{DATE}}")
    doc.add_paragraph("")
    _add_signature(doc)

    path = TEMPLATE_DIR / "demand_promissory_note.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")


def create_lc_master_agreement():
    doc = Document()
    _add_header(doc, "LC Master Agreement")
    _add_address_block(doc)

    doc.add_paragraph(
        "This Master Letter of Credit Agreement is entered into between "
        "Meezan Bank Limited ('the Bank') and {{CUSTOMER_NAME}} ('the Customer') "
        "for the issuance of Letters of Credit up to a limit of "
        "{{CURRENCY}} {{APPROVED_LIMIT}} millions."
    )
    doc.add_paragraph("")
    for heading, body in [
        ("LC Limit", "{{CURRENCY}} {{APPROVED_LIMIT}} millions"),
        ("Tenor", "{{TENOR}}"),
        ("Commission / Charges", "{{PROFIT_RATE}}"),
        ("Security", "{{SECURITY}}"),
        ("Governing Law", "The laws of Pakistan"),
    ]:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(body)

    _add_signature(doc)
    path = TEMPLATE_DIR / "lc_master_agreement.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")


def create_letter_of_credit_application():
    doc = Document()
    _add_header(doc, "Letter of Credit Application")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Details"
    for l, v in [
        ("Applicant", "{{CUSTOMER_NAME}}"),
        ("Address", "{{CUSTOMER_LOCATION}}"),
        ("LC Type", "{{NATURE_OF_LIMIT}}"),
        ("Amount", "{{CURRENCY}} {{APPROVED_LIMIT}}"),
        ("Tenor", "{{TENOR}}"),
        ("Purpose", "{{PURPOSE}}"),
        ("Commission", "{{PROFIT_RATE}}"),
        ("Security", "{{SECURITY}}"),
        ("Approval No.", "{{APPROVAL_NO}}"),
        ("Date", "{{DATE}}"),
    ]:
        r = table.add_row().cells
        r[0].text = l
        r[1].text = v

    doc.add_paragraph("")
    _add_signature(doc)
    path = TEMPLATE_DIR / "letter_of_credit_application.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")


def create_trust_receipt_agreement():
    doc = Document()
    _add_header(doc, "Trust Receipt Agreement")
    _add_address_block(doc)

    doc.add_paragraph(
        "I/We, {{CUSTOMER_NAME}}, acknowledge receipt of the documents listed below "
        "from Meezan Bank Limited relating to the Letter of Credit and agree to hold "
        "the goods or proceeds on trust for the Bank."
    )
    doc.add_paragraph("")
    for l, v in [
        ("LC / Facility Reference", "{{APPROVAL_NO}}"),
        ("Amount", "{{CURRENCY}} {{APPROVED_LIMIT}}"),
        ("Security", "{{SECURITY}}"),
        ("Date", "{{DATE}}"),
    ]:
        doc.add_paragraph(f"{l}: {v}")

    _add_signature(doc)
    path = TEMPLATE_DIR / "trust_receipt_agreement.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")


def create_msfa_agreement():
    doc = Document()
    _add_header(doc, "Murabaha Sub-Finance Agreement (MSFA)")
    _add_address_block(doc)

    doc.add_paragraph(
        "This Murabaha Sub-Finance Agreement is entered into between Meezan Bank "
        "Limited and {{CUSTOMER_NAME}} for financing of imports under the approved "
        "LC facility of {{CURRENCY}} {{APPROVED_LIMIT}} millions."
    )
    doc.add_paragraph("")
    for l, v in [
        ("Facility Type", "{{FACILITY_TYPE}}"),
        ("Limit", "{{CURRENCY}} {{APPROVED_LIMIT}} millions"),
        ("Profit Rate", "{{PROFIT_RATE}}"),
        ("Tenor", "{{TENOR}}"),
        ("Security", "{{SECURITY}}"),
    ]:
        doc.add_paragraph(f"{l}: {v}")

    _add_signature(doc)
    path = TEMPLATE_DIR / "msfa_agreement.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")


def create_letter_of_lien_and_setoff():
    doc = Document()
    _add_header(doc, "Letter of Lien and Set-Off")
    _add_address_block(doc)

    doc.add_paragraph(
        "I/We, {{CUSTOMER_NAME}}, hereby grant Meezan Bank Limited a lien over and "
        "right of set-off against all deposits, accounts and assets held with the "
        "Bank as security for the facility approved vide {{APPROVAL_NO}}."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Facility Amount: {{CURRENCY}} {{APPROVED_LIMIT}} millions")
    doc.add_paragraph("Security Details: {{SECURITY}}")
    doc.add_paragraph("Date: {{DATE}}")
    doc.add_paragraph("")
    _add_signature(doc)

    path = TEMPLATE_DIR / "letter_of_lien_and_set-off.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")


def create_personal_guarantee():
    doc = Document()
    _add_header(doc, "Personal Guarantee")
    _add_address_block(doc)

    doc.add_paragraph(
        "I/We, the undersigned Director(s) / Guarantor(s) of {{CUSTOMER_NAME}}, "
        "hereby unconditionally and irrevocably guarantee to Meezan Bank Limited "
        "the due payment of all sums owed by {{CUSTOMER_NAME}} under the facility "
        "approved vide {{APPROVAL_NO}} up to {{CURRENCY}} {{APPROVED_LIMIT}} millions."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Date: {{DATE}}")
    doc.add_paragraph("Place: {{CUSTOMER_LOCATION}}")
    doc.add_paragraph("")
    _add_signature(doc)

    path = TEMPLATE_DIR / "personal_guarantee.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")


def create_cash_margin_agreement():
    doc = Document()
    _add_header(doc, "Cash Margin Agreement")
    _add_address_block(doc)

    doc.add_paragraph(
        "{{CUSTOMER_NAME}} agrees to maintain a cash margin / lien over deposits "
        "with Meezan Bank Limited as security for the credit facility approved "
        "vide {{APPROVAL_NO}}."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Facility Amount: {{CURRENCY}} {{APPROVED_LIMIT}} millions")
    doc.add_paragraph("Security: {{SECURITY}}")
    doc.add_paragraph("Date: {{DATE}}")
    doc.add_paragraph("")
    _add_signature(doc)

    path = TEMPLATE_DIR / "cash_margin_agreement.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")


def create_master_murabaha_agreement():
    doc = Document()
    _add_header(doc, "Master Murabaha Agreement")
    _add_address_block(doc)

    doc.add_paragraph(
        "This Master Murabaha Agreement is entered into between Meezan Bank Limited "
        "('the Bank') and {{CUSTOMER_NAME}} ('the Customer') for the provision of "
        "Murabaha financing up to {{CURRENCY}} {{APPROVED_LIMIT}} millions."
    )
    doc.add_paragraph("")
    for l, v in [
        ("Facility Limit", "{{CURRENCY}} {{APPROVED_LIMIT}} millions"),
        ("Profit Markup", "{{PROFIT_RATE}}"),
        ("Tenor", "{{TENOR}}"),
        ("Purpose", "{{PURPOSE}}"),
        ("Security", "{{SECURITY}}"),
    ]:
        doc.add_paragraph(f"{l}: {v}")

    _add_signature(doc)
    path = TEMPLATE_DIR / "master_murabaha_agreement.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")


# Map of all templates to create
TEMPLATES = {
    "offer_letter": create_offer_letter,
    "sanction_letter": create_sanction_letter,
    "terms_and_conditions_sheet": create_terms_conditions_sheet,
    "demand_promissory_note": create_demand_promissory_note,
    "lc_master_agreement": create_lc_master_agreement,
    "letter_of_credit_application": create_letter_of_credit_application,
    "trust_receipt_agreement": create_trust_receipt_agreement,
    "msfa_agreement": create_msfa_agreement,
    "letter_of_lien_and_set-off": create_letter_of_lien_and_setoff,
    "personal_guarantee": create_personal_guarantee,
    "cash_margin_agreement": create_cash_margin_agreement,
    "master_murabaha_agreement": create_master_murabaha_agreement,
}


def main():
    print("=" * 60)
    print("CREATING DOCUMENT TEMPLATES")
    print(f"Output directory: {TEMPLATE_DIR}")
    print("=" * 60)
    print()

    for name, fn in TEMPLATES.items():
        try:
            fn()
        except Exception as exc:
            print(f"  ✗ {name}.docx — ERROR: {exc}")

    print()
    print(f"Done. {len(TEMPLATES)} templates written to {TEMPLATE_DIR}")
    print()
    print("Next steps:")
    print("  1. Open each .docx in Microsoft Word.")
    print("  2. Add your bank logo, fonts, and brand colours.")
    print("  3. The {{TOKENS}} will be replaced automatically at generation time.")
    print("  4. Do NOT rename or remove the tokens — the generator needs them.")


if __name__ == "__main__":
    main()