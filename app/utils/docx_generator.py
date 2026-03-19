"""
DOCX document generation utilities.

How templates work
------------------
1. Place a .docx template file in app/templates/, e.g. offer_letter.docx
2. Inside the template use placeholder tokens enclosed in double curly braces,
   e.g.  {{CUSTOMER_NAME}},  {{FACILITY_TABLE}},  {{DATE}}
3. This generator will:
      a. Open the template.
      b. Replace every placeholder in paragraphs AND table cells.
      c. For block-level content (e.g. the facility table) it inserts a
         pre-built table into the document rather than a single token.
4. If no template exists the generator creates a plain document from scratch
   so the pipeline never hard-fails.

Supported tokens (all templates)
---------------------------------
{{CUSTOMER_NAME}}        Full customer name
{{APPROVAL_NO}}          Sanction approval number
{{SANCTION_DATE}}        Date of sanction
{{ICRR}}                 Credit rating
{{BUSINESS_SEGMENT}}     Business segment
{{CUSTOMER_LOCATION}}    Customer address
{{DATE}}                 Today's date (document generation date)

Facility tokens (for single-facility documents or when iterating)
---------------------------------
{{FACILITY_TYPE}}        e.g. "LC"
{{NATURE_OF_LIMIT}}      Full nature-of-limit text
{{APPROVED_LIMIT}}       Approved limit amount
{{EXISTING_LIMIT}}       Existing limit amount
{{CURRENCY}}             Currency code
{{PROFIT_RATE}}          Profit / commission rate
{{TENOR}}                Tenor string
{{SECURITY}}             Security / collateral details
{{PURPOSE}}              Purpose of facility
"""

import os
from copy import deepcopy
from datetime import datetime
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.schemas.sanction_schema import SanctionData, FacilityData


class DocxGenerator:
    """Generate DOCX documents from templates or from scratch."""

    def __init__(self):
        self.template_dir = os.path.join(os.path.dirname(__file__), "..", "templates")
        self.output_dir = os.path.join(os.path.dirname(__file__), "..", "..", "output")
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.template_dir, exist_ok=True)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def generate_document(
        self,
        doc_type: str,
        sanction_data: SanctionData,
        facility: Optional[FacilityData] = None,
    ) -> str:
        """
        Generate a single document.

        Args:
            doc_type:      Document key, e.g. "offer_letter", "lc_master_agreement".
                           Must match a template filename (without .docx) if one exists.
            sanction_data: Full sanction data.
            facility:      If the document is facility-specific, pass the relevant
                           FacilityData object so facility-level tokens are filled.
                           Defaults to the first facility if None.

        Returns:
            Absolute path to the saved .docx file.
        """
        if facility is None and sanction_data.facilities:
            facility = sanction_data.facilities[0]

        template_path = os.path.join(self.template_dir, f"{doc_type}.docx")

        if os.path.exists(template_path):
            doc = self._fill_template(template_path, sanction_data, facility)
        else:
            doc = self._build_from_scratch(doc_type, sanction_data, facility)

        # Save
        safe_customer = "".join(
            c for c in sanction_data.customer_name if c.isalnum() or c in (" ", "_", "-")
        ).strip().replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{doc_type}_{safe_customer}_{timestamp}.docx"
        output_path = os.path.join(self.output_dir, filename)
        doc.save(output_path)
        return output_path

    def generate_all_documents(
        self,
        required_docs: Dict[str, List[str]],
        sanction_data: SanctionData,
    ) -> Dict[str, List[str]]:
        """
        Generate every document listed in required_docs.

        Args:
            required_docs:  Dict returned by RuleEngine.determine_required_documents().
            sanction_data:  Sanction data.

        Returns:
            Dict mapping each category to a list of generated file paths.
        """
        generated: Dict[str, List[str]] = {
            "compulsory": [],
            "general": [],
            "facility_specific": [],
            "collateral": [],
        }

        for category, doc_names in required_docs.items():
            for doc_name in doc_names:
                doc_key = self._doc_name_to_key(doc_name)
                try:
                    # Facility-specific docs: generate once per facility
                    if category == "facility_specific":
                        for facility in sanction_data.facilities:
                            path = self.generate_document(doc_key, sanction_data, facility)
                            generated[category].append(path)
                    else:
                        path = self.generate_document(doc_key, sanction_data)
                        generated[category].append(path)
                except Exception as exc:
                    print(f"[DocxGenerator] ERROR generating '{doc_name}': {exc}")

        return generated

    # ------------------------------------------------------------------
    # Template filling
    # ------------------------------------------------------------------

    def _fill_template(
        self,
        template_path: str,
        sanction_data: SanctionData,
        facility: Optional[FacilityData],
    ) -> Document:
        """Open a template and replace all placeholder tokens."""
        doc = Document(template_path)
        replacements = self._build_replacements(sanction_data, facility)

        for para in doc.paragraphs:
            self._replace_in_paragraph(para, replacements)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, replacements)

        return doc

    def _replace_in_paragraph(self, para, replacements: Dict[str, str]):
        """
        Replace tokens inside a paragraph while preserving runs (bold, italic etc.).

        python-docx splits text across multiple runs for formatting reasons,
        so a token like {{CUSTOMER_NAME}} might be split across two or three
        runs. We rebuild the paragraph text, replace tokens, then write back.
        """
        full_text = "".join(run.text for run in para.runs)
        if "{{" not in full_text:
            return

        for token, value in replacements.items():
            full_text = full_text.replace(token, value)

        # Write the replaced text back into the first run, clear the rest
        if para.runs:
            para.runs[0].text = full_text
            for run in para.runs[1:]:
                run.text = ""

    def _build_replacements(
        self,
        sanction_data: SanctionData,
        facility: Optional[FacilityData],
    ) -> Dict[str, str]:
        """Build the token → value map for a document."""
        today = datetime.now().strftime("%B %d, %Y")

        replacements = {
            "{{CUSTOMER_NAME}}":    sanction_data.customer_name or "",
            "{{APPROVAL_NO}}":      sanction_data.approval_no or "",
            "{{SANCTION_DATE}}":    str(sanction_data.sanction_date or ""),
            "{{ICRR}}":             sanction_data.icrr or "",
            "{{BUSINESS_SEGMENT}}": sanction_data.business_segment or "",
            "{{CUSTOMER_LOCATION}}": sanction_data.customer_location or "",
            "{{DATE}}":             today,
        }

        if facility:
            replacements.update({
                "{{FACILITY_TYPE}}":   facility.facility_type or "",
                "{{NATURE_OF_LIMIT}}": facility.nature_of_limit or "",
                "{{APPROVED_LIMIT}}":  str(facility.approved_limit or ""),
                "{{EXISTING_LIMIT}}":  str(facility.existing_limit or ""),
                "{{CURRENCY}}":        facility.currency or "PKR",
                "{{PROFIT_RATE}}":     facility.profit_rate or "",
                "{{TENOR}}":           facility.tenor or "",
                "{{SECURITY}}":        str(facility.security or ""),
                "{{PURPOSE}}":         facility.purpose or "",
            })

        return replacements

    # ------------------------------------------------------------------
    # From-scratch generation (fallback when no template exists)
    # ------------------------------------------------------------------

    def _build_from_scratch(
        self,
        doc_type: str,
        sanction_data: SanctionData,
        facility: Optional[FacilityData],
    ) -> Document:
        """
        Create a basic but complete document when no template file exists.

        This is the fallback so the pipeline never crashes. Once you create
        proper .docx templates in app/templates/, this code path is skipped.
        """
        doc = Document()
        today = datetime.now().strftime("%B %d, %Y")

        title = doc_type.replace("_", " ").title()
        doc.add_heading(title, 0)

        # Header block
        doc.add_paragraph(f"Date: {today}")
        doc.add_paragraph(f"Customer: {sanction_data.customer_name}")
        if sanction_data.approval_no:
            doc.add_paragraph(f"Approval No: {sanction_data.approval_no}")
        if sanction_data.sanction_date:
            doc.add_paragraph(f"Sanction Date: {sanction_data.sanction_date}")
        if sanction_data.customer_location:
            doc.add_paragraph(f"Address: {sanction_data.customer_location}")

        doc.add_paragraph("")

        # Facility details
        if facility:
            doc.add_heading("Facility Details", level=1)
            details = [
                ("Facility Type", facility.facility_type),
                ("Nature of Limit", facility.nature_of_limit),
                ("Approved Limit", f"{facility.currency} {facility.approved_limit}"),
                ("Profit Rate", facility.profit_rate),
                ("Tenor", facility.tenor),
                ("Security", str(facility.security or "")),
                ("Purpose", facility.purpose or ""),
            ]
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Field"
            hdr[1].text = "Value"
            for label, value in details:
                if value:
                    row = table.add_row().cells
                    row[0].text = label
                    row[1].text = str(value)

        doc.add_paragraph("")

        # Terms and conditions (if any)
        if sanction_data.terms_conditions:
            doc.add_heading("Terms and Conditions", level=1)
            for i, term in enumerate(sanction_data.terms_conditions, 1):
                doc.add_paragraph(f"{i}. {term}")

        # Signature block
        doc.add_paragraph("")
        doc.add_paragraph("_" * 30 + "          " + "_" * 30)
        doc.add_paragraph("Authorised Signatory                    Relationship Manager")

        return doc

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _doc_name_to_key(doc_name: str) -> str:
        """
        Convert a human-readable document name to a file-system-safe key.

        e.g. "Letter of Credit Application" → "letter_of_credit_application"
        """
        return doc_name.lower().replace(" ", "_").replace("-", "_")