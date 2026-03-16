"""
DOCX document parsing utilities
"""
from docx import Document
from typing import Dict, List
import os
import json
from pathlib import Path
from datetime import datetime

class DocxParser:
    """Extract text and structured data from DOCX files"""
    
    def __init__(self):
        pass
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract all text from a DOCX file (paragraphs and tables)
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text as string
        """
        doc = Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():  # Only add non-empty paragraphs
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    
    def extract_and_save_text(self, file_path: str, output_dir: str = None) -> tuple[str, str]:
        """
        Extract text from DOCX and save it to a file for verification
        
        Args:
            file_path: Path to the DOCX file
            output_dir: Directory to save extracted text (defaults to app/extracted_text)
            
        Returns:
            Tuple of (extracted_text, saved_file_path)
        """
        # Extract text
        extracted_text = self.extract_text(file_path)
        
        # Determine output directory
        if output_dir is None:
            # Default to app/extracted_text
            current_dir = Path(__file__).parent.parent
            output_dir = current_dir / "extracted_text"
        else:
            output_dir = Path(output_dir)
        
        # Create directory if it doesn't exist
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate filename with timestamp
        input_filename = Path(file_path).stem
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{input_filename}_extracted_{timestamp}.txt"
        output_path = output_dir / output_filename
        
        # Save extracted text
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"Extracted from: {file_path}\n")
            f.write(f"Extraction time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("=" * 80 + "\n\n")
            f.write(extracted_text)
        
        return extracted_text, str(output_path)
    
    def extract_from_bytes_and_save(self, docx_content: bytes, original_filename: str, output_dir: str = None) -> tuple[str, str]:
        """
        Extract text from DOCX bytes and save it to a file for verification
        
        Args:
            docx_content: DOCX file content as bytes
            original_filename: Original filename (for naming the output file)
            output_dir: Directory to save extracted text (defaults to app/extracted_text)
            
        Returns:
            Tuple of (extracted_text, saved_file_path)
        """
        # Extract text from bytes
        extracted_text = self.extract_from_bytes(docx_content)
        
        # Determine output directory
        if output_dir is None:
            # Default to app/extracted_text
            current_dir = Path(__file__).parent.parent
            output_dir = current_dir / "extracted_text"
        else:
            output_dir = Path(output_dir)
        
        # Create directory if it doesn't exist
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate filename with timestamp
        input_filename = Path(original_filename).stem
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{input_filename}_extracted_{timestamp}.txt"
        output_path = output_dir / output_filename
        
        # Save extracted text
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"Extracted from: {original_filename}\n")
            f.write(f"Extraction time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("=" * 80 + "\n\n")
            f.write(extracted_text)
        
        return extracted_text, str(output_path)
    
    def extract_structured_data(self, file_path: str) -> Dict:
        """
        Extract structured data from DOCX file
        Keeps table data (Sanction Advice/Facilities) separate from paragraphs (General Conditions)
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing structured data with separated tables and paragraphs
        """
        doc = Document(file_path)
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        tables_data = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)
        
        return {
            "paragraphs": paragraphs,  # General Conditions and narrative text
            "tables": tables_data,     # Sanction Advice and facility details
            "full_text": self.extract_text(file_path)
        }
    
    def extract_and_save_structured(self, file_path: str, output_dir: str = None) -> tuple[Dict, str]:
        """
        Extract structured data from DOCX and save it as JSON for verification
        This preserves the separation between tables (facilities) and paragraphs (conditions)
        
        Args:
            file_path: Path to the DOCX file
            output_dir: Directory to save structured data (defaults to app/extracted_text)
            
        Returns:
            Tuple of (structured_data_dict, saved_file_path)
        """
        # Extract structured data
        structured_data = self.extract_structured_data(file_path)
        
        # Determine output directory
        if output_dir is None:
            current_dir = Path(__file__).parent.parent
            output_dir = current_dir / "extracted_text"
        else:
            output_dir = Path(output_dir)
        
        # Create directory if it doesn't exist
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate filename with timestamp
        input_filename = Path(file_path).stem
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{input_filename}_structured_{timestamp}.json"
        output_path = output_dir / output_filename
        
        # Save structured data as JSON
        save_data = {
            "source_file": file_path,
            "extraction_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            "data": {
                "paragraphs": structured_data["paragraphs"],
                "tables": structured_data["tables"]
            },
            "note": "Tables contain Sanction Advice/Facilities. Paragraphs contain General Conditions."
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(save_data, f, indent=2, ensure_ascii=False)
        
        return structured_data, str(output_path)
    
    def extract_and_save_all(self, file_path: str, output_dir: str = None) -> Dict[str, str]:
        """
        Extract and save both text and structured data for complete verification
        
        Args:
            file_path: Path to the DOCX file
            output_dir: Directory to save extracted data (defaults to app/extracted_text)
            
        Returns:
            Dictionary containing paths to saved files
        """
        # Save plain text extraction
        _, text_path = self.extract_and_save_text(file_path, output_dir)
        
        # Save structured extraction
        _, structured_path = self.extract_and_save_structured(file_path, output_dir)
        
        return {
            "text_file": text_path,
            "structured_file": structured_path
        }
    
    def extract_from_bytes(self, docx_content: bytes) -> str:
        """
        Extract text from DOCX file content (bytes)
        
        Args:
            docx_content: DOCX file content as bytes
            
        Returns:
            Extracted text as string
        """
        from io import BytesIO
        doc = Document(BytesIO(docx_content))
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
